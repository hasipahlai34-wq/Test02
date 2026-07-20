"""
# ============================================================
# 文档分块策略
# ← WeKnora: internal/infrastructure/chunker/ — 分块算法
#   WeKnora 实现了多种分块策略:
#   - FixedSizeChunker: 固定大小分块 (带重叠)
#   - RecursiveChunker: 递归字符分块 (按段落/句子/词逐级分割)
#   - SemanticChunker: 基于 Embedding 相似度的语义分块
#   - MarkdownChunker: Markdown 结构感知分块
# ============================================================

本模块提供:
- 4 种分块策略的 Python 实现
- 统一的 `chunk_documents()` 入口
- 自动选择最佳策略 (默认语义分块)

设计要点:
- 所有策略遵循 LangChain TextSplitter 接口
- 保留 WeKnora 的核心算法思路 (重叠、结构感知、语义边界)
- 简化: 去掉 ContextHeader、ImageInfo、FAQ 等企业级特性
"""

from __future__ import annotations

import csv
import io
import logging
import re
from collections import Counter
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Optional

from langchain_core.documents import Document
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    TokenTextSplitter,
)

from config.settings import get_settings

logger = logging.getLogger(__name__)


class ChunkingStrategy(str, Enum):
    """
    分块策略枚举
    ← WeKnora: internal/infrastructure/chunker/ 各策略独立实现
    """
    FIXED_SIZE = "fixed_size"       # 固定大小分块 (简单粗暴)
    RECURSIVE = "recursive"         # 递归字符分块 (推荐通用)
    SEMANTIC = "semantic"           # 语义分块 (按段落/语义边界)
    MARKDOWN = "markdown"           # Markdown 结构感知分块


# ================================================================
# 分块策略实现
# ================================================================



@dataclass(frozen=True)
class DocumentProfile:
    """Lightweight document profile used by dynamic chunk planning."""

    file_path: str
    file_ext: str
    file_size: int
    total_chars: int
    estimated_tokens: int
    paragraph_count: int
    avg_paragraph_tokens: float
    heading_count: int
    table_like_ratio: float
    code_like_ratio: float
    row_count: int | None = None
    language: str = "unknown"


@dataclass(frozen=True)
class ChunkPlan:
    """Deterministic chunk plan; char fields are consumed by current splitters."""

    strategy: str
    target_tokens: int
    target_chars: int
    overlap_tokens: int
    overlap_chars: int
    min_tokens: int
    preserve_tables: bool
    preserve_code_blocks: bool
    reason: str


BASE_CHUNK_PLANS: dict[str, dict[str, int | str]] = {
    ".pdf": {"strategy": "pdf_structured", "tokens": 800, "overlap": 100},
    ".docx": {"strategy": "docx_structured", "tokens": 750, "overlap": 90},
    ".doc": {"strategy": "docx_structured", "tokens": 750, "overlap": 90},
    ".md": {"strategy": "markdown_enhanced", "tokens": 650, "overlap": 70},
    ".markdown": {"strategy": "markdown_enhanced", "tokens": 650, "overlap": 70},
    ".txt": {"strategy": "txt_paragraph", "tokens": 650, "overlap": 80},
    ".csv": {"strategy": "csv_rows", "tokens": 1000, "overlap": 0},
}


def _clamp(value: float, minimum: int, maximum: int) -> int:
    return int(max(minimum, min(maximum, round(value))))


def _read_text_with_fallback(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="ignore")


def _extract_profile_text(path: Path) -> tuple[str, int | None]:
    ext = path.suffix.lower()

    if ext == ".pdf":
        try:
            import fitz

            with fitz.open(str(path)) as doc:
                return "\n\n".join(page.get_text("text") for page in doc), None
        except Exception as exc:
            logger.debug("PDF profile extraction fallback: %s", exc)
            return "", None

    if ext in {".docx", ".doc"}:
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(path))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                parts.append(_format_table_as_text(table))
            return "\n\n".join(parts), None
        except Exception as exc:
            logger.debug("DOCX profile extraction fallback: %s", exc)
            return "", None

    if ext == ".csv":
        try:
            with open(path, "r", encoding="utf-8", newline="") as f:
                rows = list(csv.reader(f))
        except UnicodeDecodeError:
            with open(path, "r", encoding="gb18030", newline="") as f:
                rows = list(csv.reader(f))
        text = "\n".join(",".join(row) for row in rows)
        return text, max(0, len(rows) - 1)

    return _read_text_with_fallback(path), None


def _count_heading_lines(lines: list[str], file_ext: str) -> int:
    count = 0
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        is_markdown_heading = stripped.startswith("#")
        is_numbered_heading = bool(re.match(r"^\d+(?:\.\d+)*[\u3001.\s]+", stripped))
        is_chinese_heading = stripped.startswith("\u7b2c") and any(
            marker in stripped[:16]
            for marker in ("\u7ae0", "\u8282", "\u7bc7", "\u90e8\u5206")
        )
        if is_markdown_heading or is_numbered_heading or is_chinese_heading:
            count += 1
    if file_ext in {".md", ".markdown"}:
        return count
    return count + sum(1 for line in lines if 0 < len(line.strip()) <= 40 and line.strip().endswith(("\uff1a", ":")))


def profile_document(file_path: str) -> DocumentProfile:
    """Build a document profile; failures return a conservative profile."""
    path = Path(file_path)
    ext = path.suffix.lower()
    file_size = path.stat().st_size if path.exists() else 0

    try:
        text, row_count = _extract_profile_text(path)
    except Exception as exc:
        logger.warning("Document profiling failed (%s): %s", file_path, exc)
        text, row_count = "", None

    total_chars = len(text)
    estimated_tokens = max(1, int(total_chars / 2))
    nonempty_lines = [line.strip() for line in text.splitlines() if line.strip()]
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if len(paragraphs) <= 1 and nonempty_lines:
        paragraphs = nonempty_lines

    paragraph_count = len(paragraphs)
    avg_paragraph_tokens = (estimated_tokens / paragraph_count) if paragraph_count else float(estimated_tokens)
    heading_count = _count_heading_lines(nonempty_lines, ext)

    line_count = max(1, len(nonempty_lines))
    table_lines = sum(
        1
        for line in nonempty_lines
        if line.count("|") >= 2 or line.count(",") >= 3 or "	" in line
    )
    code_lines = sum(
        1
        for line in nonempty_lines
        if line.startswith(("```", "    ", "	"))
        or bool(re.search(r"[{};]\s*$", line))
        or line.startswith(("def ", "class ", "function ", "import ", "from "))
    )

    table_like_ratio = 1.0 if ext == ".csv" and row_count else table_lines / line_count
    code_like_ratio = code_lines / line_count

    return DocumentProfile(
        file_path=str(file_path),
        file_ext=ext,
        file_size=file_size,
        total_chars=total_chars,
        estimated_tokens=estimated_tokens,
        paragraph_count=paragraph_count,
        avg_paragraph_tokens=avg_paragraph_tokens,
        heading_count=heading_count,
        table_like_ratio=table_like_ratio,
        code_like_ratio=code_like_ratio,
        row_count=row_count,
    )


def infer_chunk_plan(profile: DocumentProfile) -> ChunkPlan:
    """Infer a dynamic chunk plan from file type, length, and structure."""
    base = BASE_CHUNK_PLANS.get(
        profile.file_ext,
        {"strategy": "txt_paragraph", "tokens": 650, "overlap": 80},
    )
    target_tokens = float(base["tokens"])
    overlap_tokens = float(base["overlap"])
    reasons = [f"base={base['strategy']}"]

    if profile.estimated_tokens < 800:
        target_tokens = max(200.0, float(profile.estimated_tokens))
        overlap_tokens = 0.0
        reasons.append("short_document_no_overlap")
    elif profile.estimated_tokens < 3000:
        target_tokens *= 0.85
        reasons.append("medium_document_smaller_chunks")
    elif profile.estimated_tokens > 50000:
        target_tokens *= 1.2
        reasons.append("large_document_larger_chunks")

    if profile.avg_paragraph_tokens < 40:
        target_tokens *= 0.9
        overlap_tokens *= 0.7
        reasons.append("short_paragraphs")
    elif profile.avg_paragraph_tokens > 220:
        target_tokens *= 1.15
        overlap_tokens *= 1.2
        reasons.append("long_paragraphs")

    if profile.heading_count >= 10:
        target_tokens *= 0.9
        reasons.append("many_headings")

    preserve_tables = profile.table_like_ratio > 0.2 or profile.file_ext == ".csv"
    if preserve_tables:
        overlap_tokens = min(overlap_tokens, 40.0)
        reasons.append("table_heavy")

    preserve_code_blocks = profile.code_like_ratio > 0.15
    if preserve_code_blocks:
        target_tokens = min(target_tokens, 700.0)
        overlap_tokens = min(overlap_tokens, 50.0)
        reasons.append("code_heavy")

    target_tokens_int = _clamp(target_tokens, 300, 1200)
    max_overlap = min(180, int(target_tokens_int * 0.2))
    overlap_tokens_int = _clamp(overlap_tokens, 0, max_overlap)

    return ChunkPlan(
        strategy=str(base["strategy"]),
        target_tokens=target_tokens_int,
        target_chars=max(400, int(target_tokens_int * 2)),
        overlap_tokens=overlap_tokens_int,
        overlap_chars=int(overlap_tokens_int * 2),
        min_tokens=100,
        preserve_tables=preserve_tables,
        preserve_code_blocks=preserve_code_blocks,
        reason=", ".join(reasons),
    )


def get_chunk_plan(file_path: str) -> ChunkPlan:
    return infer_chunk_plan(profile_document(file_path))

def chunk_fixed_size(
    documents: list[Document],
    chunk_size: int = 512,
    chunk_overlap: int = 50,
) -> list[Document]:
    """
    固定大小分块
    ← WeKnora: FixedSizeChunker — 最基础的分块方式
    按固定字符数切分，带重叠以保持上下文连续性

    Args:
        documents: LangChain Document 列表
        chunk_size: 每个分块的最大字符数
        chunk_overlap: 相邻分块之间的重叠字符数

    Returns:
        分块后的 Document 列表
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", ".", "，", ",", " ", ""],
        length_function=len,
    )

    chunks = splitter.split_documents(documents)
    logger.info(
        "固定大小分块: %d docs → %d chunks (size=%d, overlap=%d)",
        len(documents), len(chunks), chunk_size, chunk_overlap,
    )
    return chunks


def chunk_recursive(
    documents: list[Document],
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[Document]:
    """
    递归字符分块
    ← WeKnora: RecursiveChunker — 按优先级递归尝试分割符
    推荐默认策略: 平衡了简单性和效果

    分割符优先级: 双换行 → 单换行 → 中文句号 → 英文句号 → 逗号 → 空格 → 字符

    Args:
        documents: LangChain Document 列表
        chunk_size: 每个分块的最大字符数
        chunk_overlap: 重叠字符数

    Returns:
        分块后的 Document 列表
    """
    settings = get_settings()
    chunk_size = chunk_size if chunk_size is not None else settings.chunk_size
    chunk_overlap = chunk_overlap if chunk_overlap is not None else settings.chunk_overlap

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=settings.chunk_separators,
        length_function=len,
        is_separator_regex=False,
    )

    chunks = splitter.split_documents(documents)
    logger.info(
        "递归分块: %d docs → %d chunks (size=%d, overlap=%d)",
        len(documents), len(chunks), chunk_size, chunk_overlap,
    )
    return chunks


def chunk_semantic(
    documents: list[Document],
    chunk_size: int = 1000,
    chunk_overlap: int = 0,
    min_chunk_size: int = 100,
) -> list[Document]:
    """
    语义分块 (基于段落边界和 Embedding 相似度)
    ← WeKnora: SemanticChunker — 在语义边界上切分

    实现方式:
    1. 先按段落 (双换行) 分割为候选块
    2. 合并过短的块到相邻块
    3. 对过长的块递归分割

    这种策略比纯固定大小分块更好，因为:
    - 每个 chunk 是一个完整的语义单元 (段落)
    - 不会在句子中间截断
    - 适合大多数文档类型

    Args:
        documents: LangChain Document 列表
        chunk_size: 目标分块大小
        chunk_overlap: 重叠字符数 (语义分块通常不需要重叠)
        min_chunk_size: 最小块大小，小于此值的块会被合并

    Returns:
        分块后的 Document 列表
    """
    from langchain_experimental.text_splitter import SemanticChunker
    from src.models.embeddings import get_embedding_model

    embedding_model = get_embedding_model()

    try:
        # 使用 LangChain 实验性 SemanticChunker
        splitter = SemanticChunker(
            embeddings=embedding_model._model,  # type: ignore
            breakpoint_threshold_type="percentile",
            breakpoint_threshold_amount=90,  # 在 90% 分位点切割
            min_chunk_size=min_chunk_size,
        )
        chunks = splitter.split_documents(documents)
    except Exception as e:
        logger.warning("语义分块失败 (%s)，降级为递归分块", e)
        return chunk_recursive(documents, chunk_size, chunk_overlap)

    logger.info(
        "语义分块: %d docs → %d chunks (target_size=%d)",
        len(documents), len(chunks), chunk_size,
    )
    return chunks


def chunk_markdown(
    documents: list[Document],
    chunk_size: int = 1000,
) -> list[Document]:
    """
    Markdown 结构感知分块
    ← WeKnora: MarkdownChunker — 按 Markdown 标题层级分块

    保留 Markdown 的标题层级结构:
    - # 一级标题作为大段边界
    - ## 二级标题作为中段边界
    - ### 三级标题作为小段边界

    Args:
        documents: LangChain Document 列表
        chunk_size: 分块大小 (用于在标题段内进一步分割)

    Returns:
        分块后的 Document 列表
    """
    headers_to_split_on = [
        ("#", "h1"),
        ("##", "h2"),
        ("###", "h3"),
        ("####", "h4"),
    ]

    try:
        markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=headers_to_split_on,
            strip_headers=False,
        )

        # 第一步: 按标题层级分割
        all_splits: list[Document] = []
        for doc in documents:
            # MarkdownHeaderTextSplitter 需要纯文本输入
            raw_splits = markdown_splitter.split_text(doc.page_content)
            # 兼容处理: 不同版本可能返回 List[str] 或 List[Document]
            for split in raw_splits:
                if isinstance(split, str):
                    split = Document(page_content=split)
                split.metadata.update(doc.metadata)
                all_splits.append(split)

        # 第二步: 对过长的段进一步按递归方式分割
        recursive_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=50,
            separators=["\n\n", "\n", "。", ".", " ", ""],
        )

        final_chunks = recursive_splitter.split_documents(all_splits)
        logger.info(
            "Markdown 分块: %d docs → %d header-splits → %d chunks",
            len(documents), len(all_splits), len(final_chunks),
        )
        return final_chunks

    except Exception as e:
        logger.warning("Markdown 分块失败 (%s)，降级为递归分块", e)
        return chunk_recursive(documents, chunk_size)


# ================================================================
# ★ 企业级结构感知分块策略（新增）
# ================================================================


def _chunk_pdf_structured(
    file_path: str,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[Document]:
    """
    PDF 结构感知分块 — 基于字体大小自动识别标题层级

    流程:
    1. 用 pymupdf(fitz) 提取所有 text span + 字体信息
    2. 统计众数字号 = 正文字号
    3. 字号 > 正文 N pt → 标题,按字号差分为H1/H2/H3
    4. 短行 + 加粗 → 可能是标题
    5. 按标题边界切分,每节过长则递归细切
    6. metadata: heading_path / page_number / chunk_type
    """
    settings = get_settings()
    chunk_size = chunk_size if chunk_size is not None else settings.chunk_size
    chunk_overlap = chunk_overlap if chunk_overlap is not None else settings.chunk_overlap
    font_delta = settings.chunk_heading_font_delta

    try:
        import fitz
    except ImportError:
        logger.warning("pymupdf is not installed, fallback to TXT paragraph chunking")
        return _chunk_txt_paragraph(file_path, chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    doc = fitz.open(file_path)
    all_spans: list[dict] = []

    for page_num, page in enumerate(doc, 1):
        blocks = page.get_text("dict").get("blocks", [])
        for block in blocks:
            if "lines" not in block:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"].strip()
                    if not text:
                        continue
                    all_spans.append({
                        "text": text,
                        "size": round(span["size"], 1),
                        "bold": bool(int(span.get("flags", 0)) & 2**3),
                        "page": page_num,
                    })
    doc.close()

    if not all_spans:
        logger.warning("PDF 无文本内容,降级为 TXT 段落分块")
        return _chunk_txt_paragraph(file_path, chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    # 推断正文字号（众数）
    size_counts = Counter(s["size"] for s in all_spans)
    body_size = size_counts.most_common(1)[0][0]

    # 构建章节结构: [(heading_info, [text_spans]), ...]
    sections: list[tuple[dict, list[dict]]] = []
    current_heading: dict = {"text": "", "level": 0, "page": 1, "path": ""}
    current_spans: list[dict] = []
    heading_stack: list[tuple[str, int]] = []

    for span in all_spans:
        text = span["text"]
        size_diff = span["size"] - body_size
        is_heading = False
        h_level = 0

        # 标题判定规则
        if size_diff >= font_delta * 2 and len(text) < 120:
            is_heading = True
            h_level = 1
        elif size_diff >= font_delta and len(text) < 120:
            is_heading = True
            h_level = 2
        elif span["bold"] and size_diff >= 1 and len(text) < 80:
            is_heading = True
            h_level = 3

        if is_heading:
            # 保存上一节
            if current_spans or current_heading["text"]:
                sections.append((dict(current_heading), list(current_spans)))
                current_spans = []

            # 维护标题路径栈
            while heading_stack and heading_stack[-1][1] >= h_level:
                heading_stack.pop()
            heading_stack.append((text, h_level))

            current_heading = {
                "text": text,
                "level": h_level,
                "page": span["page"],
                "path": " > ".join(h[0] for h in heading_stack),
            }
        else:
            current_spans.append(span)

    # 最后一节
    if current_spans or current_heading["text"]:
        sections.append((current_heading, current_spans))

    # 构建 chunks
    chunks: list[Document] = []
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=settings.chunk_separators,
    )

    for heading, spans in sections:
        section_text = " ".join(s["text"] for s in spans) if spans else heading["text"]
        if heading["text"] and heading["text"] not in (section_text or ""):
            section_text = heading["text"] + "\n" + (section_text or "")

        if not section_text:
            continue

        pages = {s["page"] for s in spans} if spans else {heading.get("page", 1)}
        meta = {
            "heading_path": heading.get("path", ""),
            "heading_level": heading.get("level", 0),
            "page_number": min(pages),
            "chunk_type": "section",
            "source": str(file_path),
        }

        if len(section_text) <= chunk_size:
            chunks.append(Document(page_content=section_text, metadata=meta))
        else:
            sub_docs = splitter.create_documents([section_text], [meta])
            for sd in sub_docs:
                sd.metadata["chunk_type"] = "content"
            chunks.extend(sub_docs)

    logger.info("PDF 结构分块: %d sections → %d chunks (body_size=%.1fpt)", len(sections), len(chunks), body_size)
    return chunks


def _chunk_docx_structured(
    file_path: str,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[Document]:
    """
    Word 样式感知分块 — 基于 Heading 样式识别章节结构

    流程:
    1. 用 python-docx 读取段落 + 样式
    2. Heading 1/2/3/4 识别为标题边界
    3. 表格段落标记 table 类型,保持完整
    4. 按标题边界切分,每节过长则递归细切
    5. metadata: heading_path / heading_level / chunk_type
    """
    settings = get_settings()
    chunk_size = chunk_size if chunk_size is not None else settings.chunk_size
    chunk_overlap = chunk_overlap if chunk_overlap is not None else settings.chunk_overlap

    try:
        from docx import Document as DocxDocument
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        logger.warning("python-docx is not installed, fallback to TXT paragraph chunking")
        return _chunk_txt_paragraph(file_path, chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    doc = DocxDocument(file_path)

    # 识别标题样式
    heading_styles: set[str] = set()
    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name.startswith("Heading"):
            heading_styles.add(style.name)

    # 遍历段落构建章节
    sections: list[tuple[dict, list[str]]] = []
    current_heading: dict = {"text": "", "level": 0, "path": ""}
    current_texts: list[str] = []
    heading_stack: list[tuple[str, int]] = []
    in_table = False

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        # 检测表格上下文 (python-docx 的表格段落有特殊标记)
        # 检查是否在表格内: 段落的前一个元素是表格
        is_table_cell = False
        try:
            p_elem = para._element
            parent_tag = p_elem.getparent().tag if p_elem.getparent() is not None else ""
            # w:tc = table cell in Word XML namespace
            is_table_cell = "tc" in parent_tag.lower().split("}")[-1] if "}" in parent_tag else "tc" in parent_tag.lower()
        except Exception:
            pass

        if is_table_cell:
            # 表格单元格内容
            current_texts.append(f"[表格] {text}" if text else "")
            in_table = True
            continue

        if in_table and not text:
            in_table = False
            continue

        if not text:
            continue

        # 判断标题
        is_heading = False
        h_level = 0

        if style_name in heading_styles:
            is_heading = True
            try:
                h_level = int(style_name.split()[-1])
            except ValueError:
                h_level = 1
        # 也检测非标准标题: 加粗 + 短文本
        elif para.runs and para.runs[0].bold and len(text) < 100:
            is_heading = True
            h_level = 3

        if is_heading:
            if current_texts or current_heading["text"]:
                sections.append((dict(current_heading), list(current_texts)))
                current_texts = []

            while heading_stack and heading_stack[-1][1] >= h_level:
                heading_stack.pop()
            heading_stack.append((text, h_level))

            current_heading = {
                "text": text,
                "level": h_level,
                "path": " > ".join(h[0] for h in heading_stack),
            }
        else:
            current_texts.append(text)

    # 最后一节
    if current_texts or current_heading["text"]:
        sections.append((current_heading, current_texts))

    # 处理表格: 从 doc.tables 中提取并附加到最近章节
    for table in doc.tables:
        table_text = _format_table_as_text(table)
        if table_text and sections:
            sections[-1][1].append(table_text)

    # 构建 chunks
    chunks: list[Document] = []
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=settings.chunk_separators,
    )

    for heading, texts in sections:
        section_text = "\n".join(filter(None, texts))
        if not section_text:
            section_text = heading["text"]
        if heading["text"] and heading["text"] not in (section_text or ""):
            section_text = heading["text"] + "\n" + (section_text or "")
        if not section_text:
            continue

        meta = {
            "heading_path": heading.get("path", ""),
            "heading_level": heading.get("level", 0),
            "chunk_type": "section",
            "source": str(file_path),
        }

        if len(section_text) <= chunk_size:
            chunks.append(Document(page_content=section_text, metadata=meta))
        else:
            sub_docs = splitter.create_documents([section_text], [meta])
            for sd in sub_docs:
                sd.metadata["chunk_type"] = "content"
            chunks.extend(sub_docs)

    logger.info("DOCX 样式分块: %d sections → %d chunks", len(sections), len(chunks))
    return chunks


def _format_table_as_text(table) -> str:
    """将 python-docx Table 转为可读文本"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def _chunk_markdown_enhanced(
    file_path: str,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[Document]:
    """
    增强 Markdown 结构分块 — 保护代码块/表格/列表完整性

    与旧版 chunk_markdown 的区别:
    - 代码块(``` ```)不被切开
    - 表格(|...|)不被切开
    - 连续列表项(- / 1.)合并为一个 chunk
    - 过长段落按 sentence 边界切分
    """
    settings = get_settings()
    chunk_size = chunk_size if chunk_size is not None else settings.chunk_size
    chunk_overlap = chunk_overlap if chunk_overlap is not None else settings.chunk_overlap

    text = Path(file_path).read_text(encoding="utf-8")

    # 第一步: 按标题层级切分
    headers_to_split_on = [
        ("#", "h1"),
        ("##", "h2"),
        ("###", "h3"),
        ("####", "h4"),
    ]

    try:
        md_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=headers_to_split_on,
            strip_headers=False,
        )
        header_splits_raw = md_splitter.split_text(text)
        # 兼容处理: 不同版本可能返回 List[str] 或 List[Document]
        header_splits: list[Document] = []
        for item in header_splits_raw:
            if isinstance(item, str):
                header_splits.append(Document(page_content=item))
            elif hasattr(item, "page_content"):
                header_splits.append(item)
            else:
                logger.warning(
                    "Markdown split_text 返回未知类型 %s，尝试转为字符串",
                    type(item).__name__,
                )
                header_splits.append(Document(page_content=str(item)))
    except Exception:
        # 降级: 直接按 # 行切分
        header_splits = [Document(page_content=text, metadata={})]

    # 第二步: 检测每节的特殊块(代码/表格/列表),保护其完整性
    protected_sections: list[Document] = []

    for doc_section in header_splits:
        content = doc_section.page_content
        meta = dict(doc_section.metadata)

        # 提取 metadata 中的标题路径
        heading_parts = []
        for key in ["h1", "h2", "h3", "h4"]:
            val = meta.get(key, "")
            if val:
                heading_parts.append(val)
        heading_path = " > ".join(heading_parts)
        heading_level = len(heading_parts)
        meta["heading_path"] = heading_path
        meta["heading_level"] = heading_level
        meta["source"] = str(file_path)

        # 检测代码块
        has_code = "```" in content
        meta["has_code"] = has_code

        # 检测表格
        has_table = bool(re.search(r"^\|.+\|.+$", content, re.MULTILINE))
        meta["has_table"] = has_table

        if len(content) <= chunk_size:
            meta["chunk_type"] = "section"
            protected_sections.append(Document(page_content=content, metadata=meta))
        else:
            # 需要再切分,但保护代码块和表格
            sub_chunks = _split_with_protected_blocks(
                content, chunk_size, chunk_overlap, settings.chunk_separators, meta
            )
            protected_sections.extend(sub_chunks)

    logger.info("Markdown 增强分块: %d header-splits → %d chunks", len(header_splits), len(protected_sections))
    return protected_sections


def _split_with_protected_blocks(
    content: str,
    chunk_size: int,
    chunk_overlap: int,
    separators: list[str],
    base_meta: dict,
) -> list[Document]:
    """在切分长文本时保护代码块和表格不被截断"""
    # 提取受保护块
    protected_blocks: list[tuple[str, str]] = []  # [(placeholder, original), ...]
    placeholder_idx = 0

    # 保护代码块
    def _protect_code(m: re.Match) -> str:
        nonlocal placeholder_idx
        key = f"__CODE_BLOCK_{placeholder_idx}__"
        protected_blocks.append((key, m.group(0)))
        placeholder_idx += 1
        return key

    content = re.sub(r"```[\s\S]*?```", _protect_code, content)

    # 保护表格
    def _protect_table(m: re.Match) -> str:
        nonlocal placeholder_idx
        key = f"__TABLE_BLOCK_{placeholder_idx}__"
        protected_blocks.append((key, m.group(0)))
        placeholder_idx += 1
        return key

    content = re.sub(r"(?:^\|.+\|\n)+", _protect_table, content, flags=re.MULTILINE)

    # 保护连续列表
    def _protect_list(m: re.Match) -> str:
        nonlocal placeholder_idx
        key = f"__LIST_BLOCK_{placeholder_idx}__"
        protected_blocks.append((key, m.group(0)))
        placeholder_idx += 1
        return key

    content = re.sub(r"(?:^[\-\*]\s+.+\n?)+", _protect_list, content, flags=re.MULTILINE)
    content = re.sub(r"(?:^\d+\.\s+.+\n?)+", _protect_list, content, flags=re.MULTILINE)

    # 递归切分剩余文本
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=separators,
    )

    chunks = splitter.create_documents([content], [base_meta])

    # 还原受保护块
    final_chunks = []
    for chunk in chunks:
        c_text = chunk.page_content
        for key, original in protected_blocks:
            c_text = c_text.replace(key, original)
        final_chunks.append(Document(
            page_content=c_text,
            metadata=dict(chunk.metadata),
        ))

    return final_chunks


def _chunk_txt_paragraph(
    file_path: str,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[Document]:
    """
    TXT 段落感知分块 — 按段落边界(双换行)切分

    流程:
    1. 按 \\n\\n 分割为段落
    2. 每个段落尽量保持完整
    3. 过长段落按分隔符递归细切(中文句子优先)
    4. metadata: paragraph_index / chunk_type
    """
    settings = get_settings()
    chunk_size = chunk_size if chunk_size is not None else settings.chunk_size
    chunk_overlap = chunk_overlap if chunk_overlap is not None else settings.chunk_overlap

    text = Path(file_path).read_text(encoding="utf-8")
    source_name = Path(file_path).name

    # 按双换行分段落
    paragraphs = re.split(r"\n\s*\n", text)
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    if not paragraphs:
        return []

    chunks: list[Document] = []
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=settings.chunk_separators,
    )

    for i, para in enumerate(paragraphs):
        meta = {
            "paragraph_index": i,
            "chunk_type": "paragraph",
            "source": str(file_path),
        }

        if len(para) <= chunk_size:
            chunks.append(Document(page_content=para, metadata=meta))
        else:
            # 过长段落: 按分隔符递归切分
            sub_docs = splitter.create_documents([para], [meta])
            for sd in sub_docs:
                sd.metadata["chunk_type"] = "sentence_group"
            chunks.extend(sub_docs)

    logger.info("TXT 段落分块: %d paragraphs → %d chunks", len(paragraphs), len(chunks))
    return chunks


def _chunk_csv_rows(
    file_path: str,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[Document]:
    """
    CSV 行完整分块 — 每 N 行一组,每组携带表头

    流程:
    1. 读取 CSV,提取表头
    2. 估算每行平均长度,计算每组行数
    3. 每组携带完整表头行
    4. metadata: columns / row_range / total_rows / chunk_type
    """
    settings = get_settings()
    chunk_size = chunk_size if chunk_size is not None else settings.chunk_size
    source_name = Path(file_path).name

    # 读取 CSV
    with open(file_path, "r", encoding="utf-8", newline="") as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        logger.warning("CSV 无数据行")
        return []

    header = rows[0]
    data_rows = rows[1:]
    total_data_rows = len(data_rows)

    if total_data_rows == 0:
        return []

    # 估算每行平均长度
    avg_row_len = sum(len(",".join(r)) for r in data_rows[:min(20, total_data_rows)]) / min(20, total_data_rows)

    # 每组行数: chunk_size / 平均行长,最少 1 行,最多 200 行
    rows_per_chunk = max(1, min(200, int(chunk_size / max(avg_row_len, 1))))

    chunks: list[Document] = []
    header_text = " | ".join(header)

    for start in range(0, total_data_rows, rows_per_chunk):
        end = min(start + rows_per_chunk, total_data_rows)
        batch = data_rows[start:end]

        # 每条记录: 列名=值 格式
        lines = [header_text]  # 每组携带表头
        for row in batch:
            row_parts = []
            for j, val in enumerate(row):
                col_name = header[j] if j < len(header) else f"col{j}"
                row_parts.append(f"{col_name}={val}")
            lines.append(", ".join(row_parts))

        chunk_text = "\n".join(lines)
        meta = {
            "columns": header,
            "row_range": f"{start + 1}-{end}",
            "total_rows": total_data_rows,
            "chunk_type": "csv_rows",
            "source": str(file_path),
        }

        chunks.append(Document(page_content=chunk_text, metadata=meta))

    logger.info("CSV 分块: %d rows → %d chunks (%d rows/group)", total_data_rows, len(chunks), rows_per_chunk)
    return chunks


# ================================================================
# 后处理
# ================================================================


def _post_process_chunks(chunks: list[Document]) -> list[Document]:
    """合并过小的 chunk 到相邻 chunk,确保每个 chunk 有意义"""
    settings = get_settings()
    min_size = settings.chunk_min_size

    if not chunks:
        return chunks

    processed: list[Document] = []
    for chunk in chunks:
        # 类型守卫: 防御非 Document 类型进入后处理流程
        if isinstance(chunk, str):
            chunk = Document(page_content=chunk)
        elif not hasattr(chunk, "page_content"):
            logger.warning(
                "后处理跳过非 Document 类型 chunk: type=%s, preview=%s...",
                type(chunk).__name__,
                str(chunk)[:100],
            )
            continue

        content = chunk.page_content.strip()
        if len(content) < min_size and processed:
            # 合并到上一个 chunk
            processed[-1].page_content += "\n" + content
            if chunk.metadata:
                for k, v in chunk.metadata.items():
                    if k not in processed[-1].metadata:
                        processed[-1].metadata[k] = v
        elif len(content) >= min_size:
            processed.append(chunk)
        else:
            # 第一个 chunk 就太小,保留
            processed.append(chunk)

    skipped = len(chunks) - len(processed)
    if skipped:
        logger.debug("后处理合并了 %d 个小 chunk (阈值=%d字符)", skipped, min_size)

    return processed


# ================================================================
# ★ 新统一入口: 根据文件类型自动分发到最优策略
# ================================================================


def auto_chunk(file_path: str) -> list[Document]:
    """Create dynamic, structure-aware chunks for any supported document type."""
    source_name = Path(file_path).name
    plan = get_chunk_plan(file_path)

    def _legacy_chunks() -> list[Document]:
        ext = Path(file_path).suffix.lower()
        strategy_map: dict[str, tuple[str, callable]] = {
            ".pdf": ("pdf_structured", _chunk_pdf_structured),
            ".docx": ("docx_structured", _chunk_docx_structured),
            ".doc": ("docx_structured", _chunk_docx_structured),
            ".md": ("markdown_enhanced", _chunk_markdown_enhanced),
            ".markdown": ("markdown_enhanced", _chunk_markdown_enhanced),
            ".txt": ("txt_paragraph", _chunk_txt_paragraph),
            ".csv": ("csv_rows", _chunk_csv_rows),
        }
        strategy_name, chunk_func = strategy_map.get(ext, ("txt_paragraph_fallback", _chunk_txt_paragraph))
        try:
            legacy = chunk_func(file_path, chunk_size=plan.target_chars, chunk_overlap=plan.overlap_chars)
        except Exception as exc:
            logger.warning("%s strategy failed (%s), fallback to TXT paragraph chunking", strategy_name, exc)
            legacy = _chunk_txt_paragraph(file_path, chunk_size=plan.target_chars, chunk_overlap=plan.overlap_chars)
        return _post_process_chunks(legacy)

    try:
        from src.ingestion.document_structure import extract_document_structure

        structure = extract_document_structure(file_path)
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=plan.target_chars,
            chunk_overlap=plan.overlap_chars,
            separators=get_settings().chunk_separators,
        )
        chunks: list[Document] = []
        common_meta = {
            "document_id": structure.document_id,
            "parse_quality_score": structure.parse_quality_score,
            "outline_preview": structure.outline_text[:500],
            "element_count": len(structure.elements),
            "warning_count": len(structure.warnings),
            "parse_warnings": "; ".join(structure.warnings[:5]),
            "structure_file_type": structure.file_type,
        }

        outline_meta = {
            **common_meta,
            "source": str(file_path),
            "chunk_type": "document_outline",
            "element_type": "outline",
            "section_title": "Document Outline",
            "section_path": "Document Outline",
        }
        if structure.outline_text.strip():
            chunks.append(Document(page_content=structure.outline_text, metadata=outline_meta))

        for element in structure.elements:
            if not element.text.strip():
                continue
            meta = {
                **common_meta,
                **element.metadata,
                "source": element.source,
                "element_id": element.element_id,
                "element_type": element.element_type,
                "chunk_type": element.element_type,
                "section_title": element.section_title or "",
                "section_path": element.section_path or "",
                "heading_level": element.heading_level or 0,
                "page_number": element.page_number or 0,
                "row_range": element.row_range or "",
                "structure_order": element.order,
            }
            if len(element.text) > plan.target_chars and element.element_type not in {"outline", "row_group", "code"}:
                sub_docs = splitter.create_documents([element.text], [meta])
                for part_index, sub_doc in enumerate(sub_docs):
                    sub_doc.metadata["chunk_part"] = part_index
                    chunks.append(sub_doc)
            else:
                chunks.append(Document(page_content=element.text, metadata=meta))

        if not chunks:
            logger.warning("Structure extraction produced no chunks for %s; using legacy chunking", file_path)
            chunks = _legacy_chunks()
            structure = None
        else:
            logger.info(
                "Structure-aware chunking: %s -> %d chunks (quality=%.3f, elements=%d)",
                source_name,
                len(chunks),
                structure.parse_quality_score,
                len(structure.elements),
            )
    except Exception as exc:
        logger.warning("Structure-aware chunking failed for %s: %s; using legacy chunking", file_path, exc)
        chunks = _legacy_chunks()

    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = i
        chunk.metadata["chunking_mode"] = "dynamic_structure_aware"
        chunk.metadata["chunk_strategy"] = plan.strategy
        chunk.metadata["chunk_strategy_label"] = "structure_aware"
        chunk.metadata["chunk_target_tokens"] = plan.target_tokens
        chunk.metadata["chunk_target_chars"] = plan.target_chars
        chunk.metadata["chunk_overlap_tokens"] = plan.overlap_tokens
        chunk.metadata["chunk_overlap_chars"] = plan.overlap_chars
        chunk.metadata["chunk_plan_reason"] = plan.reason
        if "source" not in chunk.metadata:
            chunk.metadata["source"] = str(file_path)
        if "chunk_id" not in chunk.metadata:
            import uuid
            chunk.metadata["chunk_id"] = str(uuid.uuid4())[:8]

    logger.info(
        "Chunking complete: %s -> %d chunks (strategy=%s, target=%d chars, overlap=%d chars)",
        source_name,
        len(chunks),
        plan.strategy,
        plan.target_chars,
        plan.overlap_chars,
    )
    return chunks


def chunk_pdf(file_path: str) -> list[Document]:
    """Compatibility wrapper for PDF file chunking."""
    return _post_process_chunks(_chunk_pdf_structured(file_path))


def chunk_docx(file_path: str) -> list[Document]:
    """Compatibility wrapper for DOCX file chunking."""
    return _post_process_chunks(_chunk_docx_structured(file_path))


def chunk_txt(file_path: str) -> list[Document]:
    """Compatibility wrapper for TXT file chunking."""
    return _post_process_chunks(_chunk_txt_paragraph(file_path))


def chunk_csv(file_path: str) -> list[Document]:
    """Compatibility wrapper for CSV file chunking."""
    return _post_process_chunks(_chunk_csv_rows(file_path))


# ================================================================
# 旧统一分块入口(保留向后兼容)
# ================================================================


def chunk_documents(
    documents: list[Document],
    strategy: ChunkingStrategy = ChunkingStrategy.RECURSIVE,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
    **kwargs,
) -> list[Document]:
    """
    统一的文档分块入口
    ← WeKnora: chunker 模块的统一入口

    Args:
        documents: LangChain Document 列表
        strategy: 分块策略 (默认递归分块)
        chunk_size: 目标分块大小
        chunk_overlap: 重叠字符数
        **kwargs: 传递给各策略的额外参数

    Returns:
        分块后的 Document 列表，每个 Document 的 metadata 中会增加 chunk_index 字段
    """
    if not documents:
        logger.warning("没有文档需要分块")
        return []

    strategy_map = {
        ChunkingStrategy.FIXED_SIZE: chunk_fixed_size,
        ChunkingStrategy.RECURSIVE: chunk_recursive,
        ChunkingStrategy.SEMANTIC: chunk_semantic,
        ChunkingStrategy.MARKDOWN: chunk_markdown,
    }

    chunk_func = strategy_map.get(strategy)
    if chunk_func is None:
        raise ValueError(f"不支持的分块策略: {strategy}")

    settings = get_settings()
    chunk_size = chunk_size if chunk_size is not None else settings.chunk_size
    chunk_overlap = chunk_overlap if chunk_overlap is not None else settings.chunk_overlap

    chunks = chunk_func(documents, chunk_size=chunk_size, chunk_overlap=chunk_overlap, **kwargs)

    # 为每个 chunk 补充元数据 (← WeKnora: chunk.go 的 ChunkIndex, StartAt, EndAt 等)
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = i
        chunk.metadata["chunk_strategy"] = strategy.value

        # 生成唯一的 chunk_id (UUID 短格式)
        if "chunk_id" not in chunk.metadata:
            import uuid
            chunk.metadata["chunk_id"] = str(uuid.uuid4())[:8]

    logger.info(
        "分块完成: %d docs → %d chunks (策略=%s, size=%d, overlap=%d)",
        len(documents), len(chunks), strategy.value, chunk_size, chunk_overlap,
    )
    return chunks


def auto_chunk_legacy(
    documents: list[Document],
    file_type: str = "",
) -> list[Document]:
    """
    [DEPRECATED] 旧版 auto_chunk — 保留向后兼容

    新代码请使用 auto_chunk(file_path) 以获得结构感知分块。
    此函数仅做简单后缀判断: .md → Markdown, 其余 → 递归。

    Args:
        documents: LangChain Document 列表
        file_type: 文件扩展名 (.md / .pdf / .docx / .txt)

    Returns:
        分块后的 Document 列表
    """
    if file_type in (".md", ".markdown"):
        return chunk_documents(documents, strategy=ChunkingStrategy.MARKDOWN)
    else:
        return chunk_documents(documents, strategy=ChunkingStrategy.RECURSIVE)
