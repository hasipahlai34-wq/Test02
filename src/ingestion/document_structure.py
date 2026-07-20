
"""Generic document structure extraction for enterprise RAG ingestion.

This module intentionally avoids business-specific parsers such as ResumeParser
or CookbookParser. It extracts common structural elements: headings, sections,
lists, tables, code blocks, CSV schema summaries, and document outlines.
"""

from __future__ import annotations

import csv
import hashlib
import math
import re
import statistics
import uuid
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class DocumentElement:
    document_id: str
    source: str
    element_id: str
    element_type: str
    text: str
    order: int
    page_number: int | None = None
    section_title: str | None = None
    section_path: str | None = None
    heading_level: int | None = None
    row_range: str | None = None
    bbox: tuple[float, float, float, float] | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentStructure:
    document_id: str
    source: str
    file_type: str
    title: str | None
    elements: list[DocumentElement]
    outline_text: str
    parse_quality_score: float
    warnings: list[str]


_CJK = "\u4e00-\u9fff"
_HEADING_KEYWORD_RE = re.compile(
    r"(\u80cc\u666f|\u6280\u80fd|\u7ecf\u5386|\u7b80\u4ecb|\u6458\u8981|\u76ee\u5f55|"
    r"\u6750\u6599|\u6b65\u9aa4|\u505a\u6cd5|\u6ce8\u610f|\u7ae0\u8282|\u9879\u76ee|"
    r"\u81ea\u6211|\u8bc4\u4ef7|\u6559\u80b2|\u6761\u6b3e|\u5b9a\u4e49|\u7ed3\u8bba)"
)
_DATE_RANGE_RE = re.compile(r"\d{4}[./-]\d{1,2}\s*[-~\u2013\u2014]\s*(\d{4}[./-]\d{1,2}|\u81f3\u4eca|present)", re.I)
_NUMBERED_HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*|[IVXivx]+)[\u3001.\s]+\S+")
_CHINESE_CHAPTER_RE = re.compile(r"^\u7b2c[\u4e00-\u9fa5\d]+[\u7ae0\u8282\u7bc7\u90e8]")
_BULLET_RE = re.compile(r"^\s*([-*+]|\d+[.)]|[\u2022\u25cf\u25e6])\s+")


def stable_document_id(file_path: str) -> str:
    path = Path(file_path)
    try:
        stat = path.stat()
        raw = f"{path.name}:{stat.st_size}:{stat.st_mtime_ns}"
    except OSError:
        raw = f"{path}:{uuid.uuid4()}"
    digest = hashlib.sha1(raw.encode("utf-8", errors="ignore")).hexdigest()[:16]
    return f"doc-{digest}"


def normalize_text(text: str) -> str:
    text = (text or "").replace("\x00", "").replace("\u00a0", " ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines: list[str] = []
    for line in text.split("\n"):
        line = re.sub(r"[\t ]+", " ", line).strip()
        for _ in range(4):
            line = re.sub(fr"([{_CJK}])\s+([{_CJK}])", r"\1\2", line)
        line = re.sub(fr"([{_CJK}])\s+([\u3001\uff0c\u3002\uff1b\uff1a\uff01\uff1f])", r"\1\2", line)
        line = re.sub(fr"([\uff08\u300a])\s+([{_CJK}])", r"\1\2", line)
        lines.append(line)
    return "\n".join(lines).strip()


def _read_text_with_fallback(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="ignore")


def _nonempty_lines(text: str) -> list[str]:
    return [line.strip() for line in normalize_text(text).splitlines() if line.strip()]


def _looks_like_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.count("|") >= 2 or stripped.count(",") >= 3 or "\t" in stripped


def _looks_like_heading(line: str, next_line: str | None = None, *, markdown: bool = False) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if _DATE_RANGE_RE.search(stripped):
        return False
    if markdown and stripped.startswith("#"):
        return True
    if len(stripped) > 90:
        return False
    if _BULLET_RE.match(stripped):
        return False
    if _NUMBERED_HEADING_RE.match(stripped) or _CHINESE_CHAPTER_RE.match(stripped):
        return True
    if len(stripped) <= 48 and stripped.endswith((":", "\uff1a")):
        return True
    if len(stripped) <= 28 and _HEADING_KEYWORD_RE.search(stripped):
        return True
    if next_line and _DATE_RANGE_RE.search(next_line) and len(stripped) <= 70:
        return True
    return False


def _heading_level(line: str, next_line: str | None = None, *, markdown: bool = False) -> int:
    stripped = line.strip()
    if markdown and stripped.startswith("#"):
        return min(6, len(stripped) - len(stripped.lstrip("#")))
    if _NUMBERED_HEADING_RE.match(stripped):
        number = stripped.split()[0].strip(".\u3001")
        return min(4, number.count(".") + 1)
    if _CHINESE_CHAPTER_RE.match(stripped):
        return 1
    if next_line and _DATE_RANGE_RE.search(next_line):
        return 2
    if len(stripped) <= 28 and _HEADING_KEYWORD_RE.search(stripped):
        return 1
    return 2


def _safe_meta_value(value: Any) -> str | int | float | bool | None:
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    if isinstance(value, (list, tuple, set)):
        return ", ".join(str(v) for v in value)
    return str(value)


def _make_element(
    *,
    document_id: str,
    source: str,
    element_type: str,
    text: str,
    order: int,
    page_number: int | None = None,
    section_title: str | None = None,
    section_path: str | None = None,
    heading_level: int | None = None,
    row_range: str | None = None,
    metadata: dict[str, Any] | None = None,
) -> DocumentElement:
    clean_meta = {str(k): _safe_meta_value(v) for k, v in (metadata or {}).items() if v is not None}
    return DocumentElement(
        document_id=document_id,
        source=source,
        element_id=f"{document_id}:el-{order:05d}",
        element_type=element_type,
        text=normalize_text(text),
        order=order,
        page_number=page_number,
        section_title=section_title,
        section_path=section_path,
        heading_level=heading_level,
        row_range=row_range,
        metadata=clean_meta,
    )


def _extract_sections_from_lines(
    lines: list[str],
    *,
    document_id: str,
    source: str,
    file_type: str,
    start_order: int = 0,
    page_number: int | None = None,
    markdown: bool = False,
) -> list[DocumentElement]:
    elements: list[DocumentElement] = []
    order = start_order
    stack: list[tuple[int, str]] = []
    current_title: str | None = None
    current_path: str | None = None
    current_level: int | None = None
    current_lines: list[str] = []

    def flush_section() -> None:
        nonlocal order, current_lines
        text = "\n".join(line for line in current_lines if line.strip()).strip()
        if not text:
            current_lines = []
            return
        if current_title and text == current_title:
            current_lines = []
            return
        elements.append(_make_element(
            document_id=document_id,
            source=source,
            element_type="section" if current_title else "paragraph",
            text=text,
            order=order,
            page_number=page_number,
            section_title=current_title,
            section_path=current_path,
            heading_level=current_level,
            metadata={"file_type": file_type},
        ))
        order += 1
        current_lines = []

    def add_structural_group(element_type: str, group_lines: list[str]) -> None:
        nonlocal order
        text = "\n".join(line for line in group_lines if line.strip()).strip()
        if not text:
            return
        elements.append(_make_element(
            document_id=document_id,
            source=source,
            element_type=element_type,
            text=text,
            order=order,
            page_number=page_number,
            section_title=current_title,
            section_path=current_path,
            heading_level=current_level,
            metadata={"file_type": file_type, "line_count": len(group_lines)},
        ))
        order += 1

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        next_line = lines[i + 1].strip() if i + 1 < len(lines) else None
        if _looks_like_heading(line, next_line, markdown=markdown):
            flush_section()
            level = _heading_level(line, next_line, markdown=markdown)
            if line.endswith((":", "\uff1a")) and stack:
                previous_level, previous_title = stack[-1]
                if previous_title.endswith((":", "\uff1a")):
                    level = previous_level
                else:
                    level = min(6, previous_level + 1)
            heading_text = line.lstrip("#").strip() if markdown else line
            while stack and stack[-1][0] >= level:
                stack.pop()
            stack.append((level, heading_text))
            current_title = heading_text
            current_level = level
            current_path = " > ".join(title for _, title in stack)
            elements.append(_make_element(
                document_id=document_id,
                source=source,
                element_type="heading",
                text=heading_text,
                order=order,
                page_number=page_number,
                section_title=heading_text,
                section_path=current_path,
                heading_level=level,
                metadata={"file_type": file_type},
            ))
            order += 1
            current_lines = [heading_text]
        elif _looks_like_table_line(line):
            flush_section()
            table_lines: list[str] = []
            while i < len(lines) and _looks_like_table_line(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) > 1 or "|" in table_lines[0] or "\t" in table_lines[0]:
                add_structural_group("table", table_lines)
            else:
                current_lines.extend(table_lines)
            continue
        elif _BULLET_RE.match(line):
            flush_section()
            list_lines: list[str] = []
            while i < len(lines) and _BULLET_RE.match(lines[i].strip()):
                list_lines.append(lines[i].strip())
                i += 1
            add_structural_group("list", list_lines)
            continue
        else:
            current_lines.append(line)
        i += 1
    flush_section()
    return elements


def _build_outline(document_id: str, source: str, file_type: str, elements: list[DocumentElement], warnings: list[str]) -> str:
    heading_lines: list[str] = []
    for element in elements:
        if element.element_type == "heading":
            level = max(1, int(element.heading_level or 1))
            heading_lines.append(f"{'  ' * (level - 1)}- {element.text}")
    if not heading_lines:
        section_count = sum(1 for e in elements if e.element_type in {"section", "paragraph", "row_group"})
        heading_lines.append(f"- Document sections: {section_count}")
    table_count = sum(1 for e in elements if e.element_type == "table")
    list_count = sum(1 for e in elements if e.element_type == "list")
    summary = [
        "Document structural outline",
        f"document_id: {document_id}",
        f"source: {Path(source).name}",
        f"file_type: {file_type}",
        f"elements: {len(elements)}",
        f"tables: {table_count}",
        f"lists: {list_count}",
        "headings:",
        *heading_lines,
    ]
    if warnings:
        summary.extend(["warnings:", *[f"- {w}" for w in warnings[:5]]])
    return "\n".join(summary)


def _quality_score(raw_text: str, elements: list[DocumentElement], warnings: list[str]) -> float:
    if not raw_text.strip():
        return 0.0
    score = 0.55
    if len(raw_text) > 200:
        score += 0.15
    if any(e.element_type == "heading" for e in elements):
        score += 0.15
    if sum(1 for e in elements if e.element_type in {"section", "paragraph", "row_group"}) >= 2:
        score += 0.10
    spaced_cjk = len(re.findall(fr"[{_CJK}]\s+[{_CJK}]", raw_text))
    cjk_chars = len(re.findall(fr"[{_CJK}]", raw_text)) or 1
    if spaced_cjk / cjk_chars > 0.08:
        score -= 0.20
    score -= min(0.20, len(warnings) * 0.04)
    return max(0.0, min(1.0, round(score, 3)))


def _structure_from_text(file_path: str, document_id: str, file_type: str, text: str, warnings: list[str]) -> DocumentStructure:
    source = str(file_path)
    raw_text = text or ""
    lines = _nonempty_lines(raw_text)
    elements = _extract_sections_from_lines(lines, document_id=document_id, source=source, file_type=file_type)
    if not elements and raw_text.strip():
        elements = [_make_element(document_id=document_id, source=source, element_type="paragraph", text=raw_text, order=0, metadata={"file_type": file_type})]
    outline = _build_outline(document_id, source, file_type, elements, warnings)
    title = next((e.text for e in elements if e.element_type == "heading"), Path(file_path).stem)
    return DocumentStructure(
        document_id=document_id,
        source=source,
        file_type=file_type,
        title=title,
        elements=elements,
        outline_text=outline,
        parse_quality_score=_quality_score(raw_text, elements, warnings),
        warnings=warnings,
    )


def extract_pdf_structure(file_path: str, document_id: str | None = None) -> DocumentStructure:
    document_id = document_id or stable_document_id(file_path)
    warnings: list[str] = []
    try:
        import fitz
    except ImportError:
        warnings.append("pymupdf_unavailable")
        return _structure_from_text(file_path, document_id, "pdf", "", warnings)

    page_texts: list[str] = []
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                page_texts.append(page.get_text("text") or "")
    except Exception as exc:
        warnings.append(f"pdf_text_extract_failed:{type(exc).__name__}")
    raw_text = "\n\n".join(page_texts)
    if not raw_text.strip():
        warnings.append("no_extractable_text")
    structure = _structure_from_text(file_path, document_id, "pdf", raw_text, warnings)
    return structure


def extract_docx_structure(file_path: str, document_id: str | None = None) -> DocumentStructure:
    document_id = document_id or stable_document_id(file_path)
    warnings: list[str] = []
    parts: list[str] = []
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append(" | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells))
            if rows:
                parts.append("\n".join(rows))
    except Exception as exc:
        warnings.append(f"docx_extract_failed:{type(exc).__name__}")
    return _structure_from_text(file_path, document_id, "docx", "\n\n".join(parts), warnings)


def extract_txt_structure(file_path: str, document_id: str | None = None) -> DocumentStructure:
    document_id = document_id or stable_document_id(file_path)
    warnings: list[str] = []
    try:
        text = _read_text_with_fallback(Path(file_path))
    except Exception as exc:
        text = ""
        warnings.append(f"txt_read_failed:{type(exc).__name__}")
    return _structure_from_text(file_path, document_id, "txt", text, warnings)


def extract_markdown_structure(file_path: str, document_id: str | None = None) -> DocumentStructure:
    document_id = document_id or stable_document_id(file_path)
    warnings: list[str] = []
    try:
        text = _read_text_with_fallback(Path(file_path))
    except Exception as exc:
        text = ""
        warnings.append(f"markdown_read_failed:{type(exc).__name__}")
    text_without_code = re.sub(r"```[\s\S]*?```", "", text)
    lines = _nonempty_lines(text_without_code)
    elements = _extract_sections_from_lines(lines, document_id=document_id, source=str(file_path), file_type="markdown", markdown=True)
    # Add protected code block elements for better retrieval without splitting code semantics.
    order = len(elements)
    for match in re.finditer(r"```[\s\S]*?```", text):
        elements.append(_make_element(
            document_id=document_id,
            source=str(file_path),
            element_type="code",
            text=match.group(0),
            order=order,
            metadata={"file_type": "markdown"},
        ))
        order += 1
    outline = _build_outline(document_id, str(file_path), "markdown", elements, warnings)
    title = next((e.text for e in elements if e.element_type == "heading"), Path(file_path).stem)
    return DocumentStructure(document_id, str(file_path), "markdown", title, elements, outline, _quality_score(text, elements, warnings), warnings)


def _infer_column_type(values: list[str]) -> str:
    non_empty = [v.strip() for v in values if str(v).strip()]
    if not non_empty:
        return "empty"
    numeric = 0
    for value in non_empty:
        try:
            float(value.replace(",", ""))
            numeric += 1
        except ValueError:
            pass
    if numeric / len(non_empty) >= 0.8:
        return "numeric"
    date_like = sum(1 for v in non_empty if re.search(r"\d{4}[-/.]\d{1,2}([-/.]\d{1,2})?", v))
    if date_like / len(non_empty) >= 0.6:
        return "date"
    unique_ratio = len(set(non_empty)) / len(non_empty)
    if unique_ratio <= 0.5:
        return "categorical"
    return "text"


def extract_csv_structure(file_path: str, document_id: str | None = None) -> DocumentStructure:
    document_id = document_id or stable_document_id(file_path)
    warnings: list[str] = []
    rows: list[list[str]] = []
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding, newline="") as f:
                rows = list(csv.reader(f))
            break
        except UnicodeDecodeError:
            continue
        except Exception as exc:
            warnings.append(f"csv_read_failed:{type(exc).__name__}")
            break
    source = str(file_path)
    elements: list[DocumentElement] = []
    if not rows:
        warnings.append("empty_csv")
        outline = _build_outline(document_id, source, "csv", elements, warnings)
        return DocumentStructure(document_id, source, "csv", Path(file_path).stem, elements, outline, 0.0, warnings)

    header = rows[0]
    data_rows = rows[1:]
    row_count = len(data_rows)
    col_count = len(header)
    columns: dict[str, list[str]] = {name or f"col{i}": [] for i, name in enumerate(header)}
    for row in data_rows:
        for i, name in enumerate(header):
            columns[name or f"col{i}"].append(row[i] if i < len(row) else "")

    profile_lines = [f"CSV schema: {row_count} rows, {col_count} columns", "Columns:"]
    for name, values in columns.items():
        col_type = _infer_column_type(values)
        missing = sum(1 for v in values if not str(v).strip())
        line = f"- {name}: type={col_type}, missing={missing}/{max(1, row_count)}"
        if col_type == "numeric":
            nums = []
            for v in values:
                try:
                    nums.append(float(str(v).replace(",", "")))
                except ValueError:
                    pass
            if nums:
                line += f", min={min(nums):.4g}, max={max(nums):.4g}, avg={statistics.mean(nums):.4g}, sum={sum(nums):.4g}"
        elif values:
            top = Counter(v for v in values if str(v).strip()).most_common(5)
            if top:
                line += ", top=" + "; ".join(f"{k}({v})" for k, v in top)
        profile_lines.append(line)

    order = 0
    profile_text = "\n".join(profile_lines)
    elements.append(_make_element(
        document_id=document_id,
        source=source,
        element_type="outline",
        text=profile_text,
        order=order,
        metadata={"file_type": "csv", "row_count": row_count, "column_count": col_count, "columns": header},
    ))
    order += 1

    rows_per_group = 50
    header_text = " | ".join(header)
    for start in range(0, row_count, rows_per_group):
        end = min(start + rows_per_group, row_count)
        lines = [header_text]
        for row in data_rows[start:end]:
            parts = []
            for i, value in enumerate(row):
                col = header[i] if i < len(header) else f"col{i}"
                parts.append(f"{col}={value}")
            lines.append(", ".join(parts))
        elements.append(_make_element(
            document_id=document_id,
            source=source,
            element_type="row_group",
            text="\n".join(lines),
            order=order,
            row_range=f"{start + 1}-{end}",
            metadata={"file_type": "csv", "row_count": row_count, "column_count": col_count, "columns": header},
        ))
        order += 1

    outline = _build_outline(document_id, source, "csv", elements, warnings)
    outline = profile_text + "\n\n" + outline
    return DocumentStructure(document_id, source, "csv", Path(file_path).stem, elements, outline, _quality_score(profile_text, elements, warnings), warnings)


def extract_document_structure(file_path: str, document_id: str | None = None) -> DocumentStructure:
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf_structure(file_path, document_id)
    if ext in {".docx", ".doc"}:
        return extract_docx_structure(file_path, document_id)
    if ext in {".md", ".markdown"}:
        return extract_markdown_structure(file_path, document_id)
    if ext == ".csv":
        return extract_csv_structure(file_path, document_id)
    return extract_txt_structure(file_path, document_id)
